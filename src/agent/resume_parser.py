import logging
import os

logger = logging.getLogger(__name__)


def parse_resume(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        text = _parse_docx(path)
    else:
        # Plain text or no extension (e.g. src/agent/myresume)
        text = _parse_text(path)

    logger.info("Resume parsed: %d characters from '%s'", len(text), path)
    return text


def _parse_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    lines = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t and t not in lines:
                    lines.append(t)

    return "\n".join(lines)


def _parse_text(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="replace") as fh:
        return fh.read().strip()
